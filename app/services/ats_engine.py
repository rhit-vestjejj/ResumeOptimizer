from __future__ import annotations

import json
import math
import re
import zipfile
from copy import deepcopy
from datetime import datetime, timezone
import hashlib
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Set, Tuple

import pdfplumber
from docx import Document
from docx.shared import Pt

from app.models import (
    CanonicalResume,
    DateRange,
    EducationEntry,
    ExperienceEntry,
    Identity,
    JDAnalysis,
    ProjectEntry,
    Skills,
)
from app.services.extractors import extract_resume_text, heuristic_resume_to_canonical
from app.services.llm import LLMService
from app.storage import load_json, save_json
from app.utils import normalize_token, tokenize, unique_preserve_order

ISSUE_SCANNED_PDF = 'ATS001_SCANNED_PDF'
ISSUE_MULTI_COLUMN = 'ATS002_MULTI_COLUMN'
ISSUE_TABLE_TEXTBOX = 'ATS003_TABLE_TEXTBOX'
ISSUE_HEADER_FOOTER_CONTACT = 'ATS004_HEADER_FOOTER_CONTACT'
ISSUE_NONSTANDARD_BULLETS = 'ATS005_NONSTANDARD_BULLETS'
ISSUE_BROKEN_READING_ORDER = 'ATS006_BROKEN_READING_ORDER'
ISSUE_SENSITIVE_DATA = 'ATS007_SENSITIVE_DATA'

MONTHS = {
    'jan': 1,
    'january': 1,
    'feb': 2,
    'february': 2,
    'mar': 3,
    'march': 3,
    'apr': 4,
    'april': 4,
    'may': 5,
    'jun': 6,
    'june': 6,
    'jul': 7,
    'july': 7,
    'aug': 8,
    'august': 8,
    'sep': 9,
    'sept': 9,
    'september': 9,
    'oct': 10,
    'october': 10,
    'nov': 11,
    'november': 11,
    'dec': 12,
    'december': 12,
}

TITLE_TOKENS = {
    'engineer', 'developer', 'scientist', 'analyst', 'architect', 'manager', 'intern', 'researcher', 'consultant',
}

SECTION_EVIDENCE_WEIGHTS: Dict[str, float] = {
    'experience': 1.0,
    'projects': 0.95,
    'skills': 0.75,
    'certifications': 0.7,
    'education': 0.55,
    'summary': 0.45,
    'identity': 0.35,
}
MUST_HAVE_MIN_REQUIRED_HARD_SKILLS = 3
MUST_HAVE_REQUIRED_HARD_COVERAGE = 0.45

DEFAULT_SKILL_ALIAS_GRAPH: Dict[str, List[str]] = {
    'python': ['python', 'py'],
    'sql': ['sql', 'postgresql', 'mysql', 'sqlite'],
    'docker': ['docker', 'containerization'],
    'kubernetes': ['kubernetes', 'k8s'],
    'aws': ['aws', 'amazon web services'],
    'gcp': ['gcp', 'google cloud'],
    'fastapi': ['fastapi'],
    'pytorch': ['pytorch', 'torch'],
    'tensorflow': ['tensorflow', 'tf'],
    'sklearn': ['sklearn', 'scikit-learn', 'scikit learn'],
    'xgboost': ['xgboost'],
    'spark': ['spark', 'apache spark'],
    'airflow': ['airflow', 'apache airflow'],
    'kafka': ['kafka', 'apache kafka'],
    'ml': ['machine learning', 'ml'],
    'nlp': ['nlp', 'natural language processing'],
    'api': ['api', 'rest api', 'apis'],
    'ci_cd': ['ci/cd', 'continuous integration', 'continuous delivery'],
    'git': ['git', 'github', 'gitlab'],
}

SOFT_SKILL_ALIAS_GRAPH: Dict[str, List[str]] = {
    'communication': [
        'communication', 'communicate', 'present', 'presentation', 'written communication', 'verbal communication',
    ],
    'collaboration': [
        'collaboration', 'collaborate', 'cross-functional', 'cross functional', 'team player', 'partnered',
    ],
    'leadership': [
        'leadership', 'led', 'leading', 'mentor', 'mentored', 'ownership', 'owned',
    ],
    'problem_solving': [
        'problem solving', 'solve', 'solved', 'troubleshoot', 'debug', 'analytical', 'critical thinking',
    ],
    'stakeholder_management': [
        'stakeholder', 'stakeholders', 'client-facing', 'customer-facing', 'partner management',
    ],
    'adaptability': [
        'adaptability', 'adapt', 'fast-paced', 'ambiguity',
    ],
    'project_management': [
        'project management', 'planning', 'roadmap', 'prioritize', 'prioritization', 'coordination',
    ],
    'initiative': [
        'initiative', 'self-starter', 'self starter', 'proactive',
    ],
}

DEGREE_TOKENS = {
    'bachelor': {'bachelor', 'bs', 'b.s', 'ba', 'b.a'},
    'master': {'master', 'ms', 'm.s', 'mba', 'm.eng'},
    'phd': {'phd', 'ph.d', 'doctorate', 'doctoral'},
}


def _clean_text(value: str) -> str:
    return re.sub(r'\s+', ' ', value or '').strip()


def _normalize_canonical_id(value: Any) -> str:
    cleaned = _clean_text(str(value)).lower()
    if not cleaned:
        return ''
    cleaned = re.sub(r'[\s\-/]+', '_', cleaned)
    cleaned = re.sub(r'[^a-z0-9_+#\.]', '', cleaned)
    cleaned = re.sub(r'_+', '_', cleaned).strip('_')
    return cleaned


def _canonical_alias_graph(
    base_graph: Dict[str, List[str]],
    additional_aliases: Optional[Dict[str, List[str]]] = None,
) -> Dict[str, List[str]]:
    merged: Dict[str, Set[str]] = {}

    def add_entry(raw_canonical: Any, aliases: Sequence[Any]) -> None:
        canonical_id = _normalize_canonical_id(raw_canonical)
        if not canonical_id:
            return
        bucket = merged.setdefault(canonical_id, set())

        raw_cleaned = _clean_text(str(raw_canonical)).lower()
        if raw_cleaned:
            bucket.add(raw_cleaned)

        # Preserve canonical-token and phrase forms for robust matching.
        bucket.add(canonical_id)
        bucket.add(canonical_id.replace('_', ' '))

        for alias in aliases or []:
            alias_cleaned = _clean_text(str(alias)).lower()
            if alias_cleaned:
                bucket.add(alias_cleaned)

    for canonical, aliases in base_graph.items():
        add_entry(canonical, aliases)
    if additional_aliases:
        for canonical, aliases in additional_aliases.items():
            add_entry(canonical, aliases)

    return {canonical_id: sorted(values) for canonical_id, values in sorted(merged.items())}


def _extract_alias_matches(text: str, alias_map: Dict[str, List[str]]) -> List[Dict[str, Any]]:
    lowered = text.lower()
    matches: List[Dict[str, Any]] = []
    for canonical_id, aliases in sorted(alias_map.items()):
        normalized_aliases = sorted({alias.lower() for alias in aliases if _clean_text(alias)}, key=len, reverse=True)
        for alias in normalized_aliases:
            pattern = re.compile(rf'(?<![A-Za-z0-9]){re.escape(alias)}(?![A-Za-z0-9])', flags=re.IGNORECASE)
            for match in pattern.finditer(lowered):
                matches.append(
                    {
                        'canonical_id': canonical_id,
                        'matched_text': lowered[match.start():match.end()],
                        'evidence_span': [match.start(), match.end()],
                    }
                )
    return matches


def _canonical_skill_graph(alias_graph: Optional[Dict[str, List[str]]] = None) -> Dict[str, List[str]]:
    return _canonical_alias_graph(DEFAULT_SKILL_ALIAS_GRAPH, alias_graph)


def _canonical_soft_skill_graph(alias_graph: Optional[Dict[str, List[str]]] = None) -> Dict[str, List[str]]:
    return _canonical_alias_graph(SOFT_SKILL_ALIAS_GRAPH, alias_graph)


def _identity_from_text(lines: Sequence[str], text: str) -> Identity:
    email_match = re.search(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}', text)
    phone_match = re.search(r'(\+?\d[\d\s\-\(\)]{7,}\d)', text)
    links = re.findall(r'https?://[^\s)]+', text)

    name = lines[0].strip() if lines else 'Unknown Name'
    location = ''
    for line in lines[1:8]:
        if len(line) > 70:
            continue
        if ',' in line or re.search(r'\b[A-Z]{2}\b', line):
            location = line
            break

    return Identity(
        name=name,
        email=email_match.group(0) if email_match else '',
        phone=phone_match.group(0) if phone_match else '',
        location=location,
        links=links,
    )


def _split_sections(lines: Sequence[str]) -> Dict[str, List[str]]:
    headers = {
        'summary': {'summary', 'profile'},
        'education': {'education', 'academic'},
        'experience': {'experience', 'work experience', 'employment'},
        'projects': {'projects', 'project experience'},
        'skills': {'skills', 'technical skills', 'technologies'},
        'certifications': {'certifications', 'certificates'},
        'awards': {'awards', 'honors', 'achievements'},
    }
    output: Dict[str, List[str]] = {key: [] for key in headers}
    current: Optional[str] = None
    for line in lines:
        lowered = line.strip().lower().rstrip(':')
        detected = None
        for key, names in headers.items():
            if lowered in names:
                detected = key
                break
        if detected:
            current = detected
            continue
        if current:
            output[current].append(line)
    return output


def _section_parser_resume(raw_text: str) -> CanonicalResume:
    lines = [line.strip() for line in raw_text.splitlines() if line.strip()]
    sections = _split_sections(lines)
    identity = _identity_from_text(lines, raw_text)

    summary = None
    if sections['summary']:
        summary = _clean_text(' '.join(sections['summary'][:2]))

    education: List[EducationEntry] = []
    for line in sections['education']:
        if len(line) < 4:
            continue
        if not any(token in line.lower() for token in ('university', 'college', 'institute', 'school')):
            continue
        degree = 'Bachelor of Science' if 'bachelor' in line.lower() else ''
        major = ''
        if ' in ' in line.lower():
            major = line.split(' in ', 1)[1].strip()[:80]
        education.append(
            EducationEntry(
                school=line[:120],
                degree=degree,
                major=major,
                minors=[],
                gpa='',
                dates=DateRange(),
                coursework=[],
            )
        )

    experience: List[ExperienceEntry] = []
    current_exp: Optional[ExperienceEntry] = None
    for line in sections['experience']:
        bullet = re.sub(r'^[\-\*•]\s*', '', line).strip()
        if bullet != line and current_exp is not None:
            current_exp.bullets.append(bullet[:220])
            continue
        if any(token in line.lower() for token in TITLE_TOKENS):
            if current_exp:
                experience.append(current_exp)
            current_exp = ExperienceEntry(
                company='',
                title=line[:90],
                location='',
                dates=DateRange(),
                bullets=[],
            )
            continue
        if current_exp is not None and line:
            current_exp.bullets.append(line[:220])
    if current_exp:
        experience.append(current_exp)

    projects: List[ProjectEntry] = []
    current_project: Optional[ProjectEntry] = None
    for line in sections['projects']:
        bullet = re.sub(r'^[\-\*•]\s*', '', line).strip()
        if bullet != line and current_project is not None:
            current_project.bullets.append(bullet[:220])
            continue
        if len(line) <= 120:
            if current_project:
                projects.append(current_project)
            current_project = ProjectEntry(
                name=line[:120],
                tech=[],
                bullets=[],
            )
            continue
    if current_project:
        projects.append(current_project)

    skill_categories: Dict[str, List[str]] = {}
    for line in sections['skills']:
        if ':' in line:
            category, entries = line.split(':', 1)
            values = [entry.strip() for entry in entries.split(',') if entry.strip()]
            if values:
                skill_categories[_clean_text(category)] = values
    if not skill_categories and sections['skills']:
        flattened = [entry.strip() for entry in ','.join(sections['skills']).split(',') if entry.strip()]
        if flattened:
            skill_categories['General'] = flattened

    certifications = [line[:120] for line in sections['certifications'] if line.strip()]
    awards = [line[:120] for line in sections['awards'] if line.strip()]

    return CanonicalResume(
        identity=identity,
        summary=summary,
        education=education,
        experience=experience,
        projects=projects,
        skills=Skills(categories=skill_categories),
        certifications=certifications,
        awards=awards,
    )


def normalize_parses(parses: Sequence[CanonicalResume]) -> CanonicalResume:
    if not parses:
        raise ValueError('normalize_parses requires at least one parser output.')

    def pick_text(values: Iterable[str]) -> str:
        cleaned = [_clean_text(value) for value in values if _clean_text(value)]
        if not cleaned:
            return ''
        counts: Dict[str, int] = {}
        for value in cleaned:
            counts[value] = counts.get(value, 0) + 1
        ranked = sorted(counts.items(), key=lambda item: (-item[1], -len(item[0]), item[0]))
        return ranked[0][0]

    identities = [parsed.identity for parsed in parses]
    merged_identity = Identity(
        name=pick_text(identity.name for identity in identities),
        email=pick_text(identity.email for identity in identities),
        phone=pick_text(identity.phone for identity in identities),
        location=pick_text(identity.location for identity in identities),
        links=unique_preserve_order(
            link for identity in identities for link in identity.links if _clean_text(link)
        ),
    )

    preferred = sorted(
        parses,
        key=lambda parsed: (
            len(parsed.experience) + len(parsed.projects) + len(parsed.education),
            len(parsed.skills.categories),
        ),
        reverse=True,
    )[0]

    merged = CanonicalResume.model_validate(preferred.model_dump())
    merged.identity = merged_identity
    merged.summary = pick_text(parsed.summary or '' for parsed in parses) or preferred.summary
    merged.certifications = unique_preserve_order(
        cert for parsed in parses for cert in parsed.certifications if _clean_text(cert)
    )
    merged.awards = unique_preserve_order(
        award for parsed in parses for award in (parsed.awards or []) if _clean_text(award)
    )

    merged_skills: Dict[str, List[str]] = {}
    for parsed in parses:
        for category, entries in parsed.skills.categories.items():
            bucket = merged_skills.setdefault(category, [])
            for entry in entries:
                if _clean_text(entry) and entry not in bucket:
                    bucket.append(entry)
    merged.skills = Skills(categories=merged_skills)
    return merged


def compute_parse_quality(parses: Sequence[CanonicalResume], canonical: CanonicalResume) -> Dict[str, Any]:
    if not parses:
        return {
            'parse_quality': 0.0,
            'completeness_score': 0.0,
            'agreement_score': 0.0,
            'missing_critical_penalties': [],
            'details': {},
        }

    checks = {
        'name': bool(_clean_text(canonical.identity.name)),
        'email': bool(re.fullmatch(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}', canonical.identity.email)),
        'phone': bool(re.search(r'\d{10,}', re.sub(r'\D', '', canonical.identity.phone))),
        'location': bool(_clean_text(canonical.identity.location)),
        'education': len(canonical.education) > 0,
        'experience_or_project': len(canonical.experience) > 0 or len(canonical.projects) > 0,
        'skills': any(canonical.skills.categories.values()),
    }
    completeness_score = (sum(1 for passed in checks.values() if passed) / len(checks)) * 100.0

    identity_fields = ['name', 'email', 'phone', 'location']
    agreement_values: List[float] = []
    for field in identity_fields:
        field_values = [_clean_text(getattr(parsed.identity, field)) for parsed in parses]
        field_values = [value for value in field_values if value]
        if not field_values:
            continue
        most_common = max(set(field_values), key=field_values.count)
        agreement_values.append(field_values.count(most_common) / len(field_values))

    list_agreements: List[float] = []
    for field in ('experience', 'projects', 'education'):
        sizes = [len(getattr(parsed, field)) for parsed in parses]
        if not sizes:
            continue
        if max(sizes) == 0:
            list_agreements.append(1.0)
            continue
        list_agreements.append(min(sizes) / max(sizes))
    agreement_values.extend(list_agreements)
    agreement_score = (sum(agreement_values) / len(agreement_values) * 100.0) if agreement_values else 0.0

    penalties: List[Dict[str, Any]] = []
    for key, passed in checks.items():
        if passed:
            continue
        penalty = 10.0 if key in {'name', 'email', 'phone', 'experience_or_project'} else 6.0
        penalties.append({'field': key, 'penalty': penalty})
    penalty_total = sum(row['penalty'] for row in penalties)

    parse_quality = max(0.0, min(100.0, (completeness_score * 0.62) + (agreement_score * 0.38) - penalty_total))
    return {
        'parse_quality': round(parse_quality, 2),
        'completeness_score': round(completeness_score, 2),
        'agreement_score': round(agreement_score, 2),
        'missing_critical_penalties': penalties,
        'details': checks,
    }


def parse_mirror(raw_text: str, llm: Optional[LLMService] = None) -> Dict[str, Any]:
    parser_outputs: List[Tuple[str, CanonicalResume]] = [
        ('heuristic_parser', heuristic_resume_to_canonical(raw_text)),
        ('section_parser', _section_parser_resume(raw_text)),
    ]
    if llm and llm.available:
        try:
            parser_outputs.append(('llm_parser', llm.extract_canonical_resume(raw_text)))
        except Exception:
            # Keep deterministic parser pair as minimum viable mirror.
            pass

    normalized = normalize_parses([parsed for _, parsed in parser_outputs])
    quality = compute_parse_quality([parsed for _, parsed in parser_outputs], normalized)

    return {
        'parsers': [
            {'name': name, 'resume': parsed.model_dump(mode='json')}
            for name, parsed in parser_outputs
        ],
        'canonical': normalized.model_dump(mode='json'),
        'quality': quality,
    }


def score_parse_quality(parse_mirror_result: Dict[str, Any]) -> Dict[str, Any]:
    quality = parse_mirror_result.get('quality')
    if isinstance(quality, dict):
        return quality
    parser_rows = parse_mirror_result.get('parsers', [])
    parses = [CanonicalResume.model_validate(row['resume']) for row in parser_rows]
    canonical = CanonicalResume.model_validate(parse_mirror_result.get('canonical', {}))
    return compute_parse_quality(parses, canonical)


def _issue(code: str, severity: str, message: str, recommendation: str, evidence: str = '') -> Dict[str, Any]:
    return {
        'code': code,
        'severity': severity,
        'message': message,
        'recommendation': recommendation,
        'evidence': evidence,
    }


def _lint_pdf(path: Path) -> List[Dict[str, Any]]:
    issues: List[Dict[str, Any]] = []
    with pdfplumber.open(path) as pdf:
        pages = pdf.pages
        all_text = '\n'.join((page.extract_text() or '') for page in pages)
        image_pages = sum(1 for page in pages if page.images)
        if len(_clean_text(all_text)) < 80 and image_pages > 0:
            issues.append(
                _issue(
                    ISSUE_SCANNED_PDF,
                    'high',
                    'PDF appears image-based with minimal extractable text.',
                    'Export a text-based PDF directly from DOCX/LaTeX instead of scanning.',
                    evidence=f'pages={len(pages)}, image_pages={image_pages}',
                )
            )

        table_pages = 0
        nonstandard_bullets = 0
        column_hits = 0
        reading_order_pages = 0
        header_footer_contact_pages = 0

        contact_pattern = re.compile(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}|(\+?\d[\d\-\s\(\)]{8,}\d)')
        weird_bullets = re.compile(r'[▪◦►●■◆❖➤]')

        for page in pages:
            if page.extract_tables():
                table_pages += 1

            text = page.extract_text() or ''
            nonstandard_bullets += len(weird_bullets.findall(text))

            words = page.extract_words() or []
            if len(words) > 80:
                left = sum(1 for word in words if float(word.get('x0', 0.0)) < 220.0)
                right = sum(1 for word in words if float(word.get('x0', 0.0)) > 320.0)
                if left > 20 and right > 20:
                    column_hits += 1

                tops = [float(word.get('top', 0.0)) for word in words]
                back_jumps = sum(1 for idx in range(1, len(tops)) if tops[idx] + 4 < tops[idx - 1])
                if back_jumps / max(1, len(tops)) > 0.2:
                    reading_order_pages += 1

            lines = [line for line in text.splitlines() if line.strip()]
            if not lines:
                continue
            contact_lines = [idx for idx, line in enumerate(lines) if contact_pattern.search(line)]
            if contact_lines and all(idx in {0, len(lines) - 1} for idx in contact_lines):
                header_footer_contact_pages += 1

        if column_hits > 0:
            issues.append(
                _issue(
                    ISSUE_MULTI_COLUMN,
                    'high',
                    'Potential multi-column layout detected.',
                    'Use a single-column resume layout for ATS parsing reliability.',
                    evidence=f'pages_with_column_signals={column_hits}',
                )
            )
        if table_pages > 0:
            issues.append(
                _issue(
                    ISSUE_TABLE_TEXTBOX,
                    'high',
                    'Tables detected in the PDF.',
                    'Replace tables/text boxes with plain section headers and bullets.',
                    evidence=f'pages_with_tables={table_pages}',
                )
            )
        if header_footer_contact_pages > 0:
            issues.append(
                _issue(
                    ISSUE_HEADER_FOOTER_CONTACT,
                    'medium',
                    'Contact info appears isolated in header/footer lines.',
                    'Move contact details into the main body top line.',
                    evidence=f'pages_with_header_footer_contact={header_footer_contact_pages}',
                )
            )
        if nonstandard_bullets > 0:
            issues.append(
                _issue(
                    ISSUE_NONSTANDARD_BULLETS,
                    'medium',
                    'Non-standard bullet glyphs detected.',
                    'Use standard bullets (•) or hyphen bullets (-).',
                    evidence=f'nonstandard_bullet_count={nonstandard_bullets}',
                )
            )
        if reading_order_pages > 0:
            issues.append(
                _issue(
                    ISSUE_BROKEN_READING_ORDER,
                    'medium',
                    'Potential broken reading order detected.',
                    'Flatten layout and avoid floating blocks that disrupt text flow.',
                    evidence=f'pages_with_reading_order_signals={reading_order_pages}',
                )
            )
    return issues


def _lint_docx(path: Path) -> List[Dict[str, Any]]:
    issues: List[Dict[str, Any]] = []
    doc = Document(str(path))

    if doc.tables:
        issues.append(
            _issue(
                ISSUE_TABLE_TEXTBOX,
                'high',
                'DOCX contains tables.',
                'Replace tables/text boxes with plain text sections.',
                evidence=f'table_count={len(doc.tables)}',
            )
        )

    weird_bullets = re.compile(r'^[\s]*(?:▪|◦|►|●|■|◆|❖|➤)')
    weird_count = sum(1 for paragraph in doc.paragraphs if weird_bullets.search(paragraph.text or ''))
    if weird_count:
        issues.append(
            _issue(
                ISSUE_NONSTANDARD_BULLETS,
                'medium',
                'DOCX uses non-standard bullet glyphs.',
                'Replace with standard bullet or hyphen bullets.',
                evidence=f'paragraphs_with_nonstandard_bullets={weird_count}',
            )
        )

    contact_pattern = re.compile(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}|(\+?\d[\d\-\s\(\)]{8,}\d)')
    header_footer_hits = 0
    for section in doc.sections:
        header_text = '\n'.join(paragraph.text for paragraph in section.header.paragraphs)
        footer_text = '\n'.join(paragraph.text for paragraph in section.footer.paragraphs)
        if contact_pattern.search(header_text) or contact_pattern.search(footer_text):
            header_footer_hits += 1
    if header_footer_hits:
        issues.append(
            _issue(
                ISSUE_HEADER_FOOTER_CONTACT,
                'medium',
                'Contact details detected in DOCX header/footer.',
                'Move contact info into the body at the top of page one.',
                evidence=f'sections_with_header_footer_contact={header_footer_hits}',
            )
        )

    xml = doc._element.xml  # pylint: disable=protected-access
    if 'w:num="2"' in xml or 'w:num="3"' in xml:
        issues.append(
            _issue(
                ISSUE_MULTI_COLUMN,
                'high',
                'DOCX appears to use multi-column section settings.',
                'Use single-column sections for ATS compatibility.',
                evidence='section_columns>1',
            )
        )
    if 'txbxContent' in xml:
        issues.append(
            _issue(
                ISSUE_TABLE_TEXTBOX,
                'high',
                'Text boxes detected in DOCX.',
                'Move text from text boxes into normal paragraphs.',
                evidence='txbxContent',
            )
        )
    return issues


def lint_resume(path: Path) -> Dict[str, Any]:
    suffix = path.suffix.lower()
    issues: List[Dict[str, Any]] = []
    if suffix == '.pdf':
        issues.extend(_lint_pdf(path))
    elif suffix == '.docx':
        issues.extend(_lint_docx(path))
    else:
        text = path.read_text(encoding='utf-8', errors='ignore')
        weird_bullets = re.findall(r'[▪◦►●■◆❖➤]', text)
        if weird_bullets:
            issues.append(
                _issue(
                    ISSUE_NONSTANDARD_BULLETS,
                    'medium',
                    'Non-standard bullet glyphs detected in text input.',
                    'Replace with standard bullets or hyphens.',
                    evidence=f'count={len(weird_bullets)}',
                )
            )

    severity_rank = {'low': 1, 'medium': 2, 'high': 3}
    max_severity = 'none'
    if issues:
        max_severity = sorted(issues, key=lambda row: severity_rank.get(row['severity'], 0), reverse=True)[0]['severity']
    return {
        'issues': issues,
        'issue_count': len(issues),
        'max_severity': max_severity,
    }


def _is_generic_skill_category(category: str) -> bool:
    normalized = normalize_token(category or '')
    return normalized in {'skills', 'core_skills', 'technical_skills', 'technologies', 'tools'}


def render_txt(resume: CanonicalResume, output_dir: Path, filename: str = 'resume.txt') -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    lines: List[str] = []
    lines.append(resume.identity.name)
    contact = ' | '.join(
        [part for part in [resume.identity.email, resume.identity.phone, resume.identity.location, *resume.identity.links] if _clean_text(part)]
    )
    if contact:
        lines.append(contact)
    lines.append('')
    if resume.summary:
        lines.append('SUMMARY')
        lines.append(_clean_text(resume.summary))
        lines.append('')

    if resume.education:
        lines.append('EDUCATION')
        for entry in resume.education:
            date_text = ' - '.join([part for part in [entry.dates.start, entry.dates.end] if _clean_text(part)])
            header = f'{entry.school}'
            if date_text:
                header = f'{header} ({date_text})'
            lines.append(header)
            degree_line = entry.degree
            if entry.major:
                degree_line = f'{degree_line} in {entry.major}'
            if degree_line.strip():
                lines.append(degree_line)
            if entry.coursework:
                lines.append(f'Coursework: {", ".join(entry.coursework)}')
        lines.append('')

    if resume.experience:
        lines.append('EXPERIENCE')
        for entry in resume.experience:
            date_text = ' - '.join([part for part in [entry.dates.start, entry.dates.end] if _clean_text(part)])
            lines.append(f'{entry.title} | {entry.company}' + (f' ({date_text})' if date_text else ''))
            for bullet in entry.bullets:
                lines.append(f'- {bullet}')
        lines.append('')

    if resume.projects:
        lines.append('PROJECTS')
        for entry in resume.projects:
            date_text = ''
            if entry.dates:
                date_text = ' - '.join([part for part in [entry.dates.start, entry.dates.end] if _clean_text(part)])
            lines.append(entry.name + (f' ({date_text})' if date_text else ''))
            if entry.tech:
                lines.append(f'Tech: {", ".join(entry.tech)}')
            for bullet in entry.bullets:
                lines.append(f'- {bullet}')
        lines.append('')

    if resume.skills.categories:
        lines.append('SKILLS')
        for category, entries in resume.skills.categories.items():
            if _is_generic_skill_category(category):
                lines.append(f'{", ".join(entries)}')
            else:
                lines.append(f'{category}: {", ".join(entries)}')
        lines.append('')

    if resume.certifications:
        lines.append('CERTIFICATIONS')
        for cert in resume.certifications:
            lines.append(f'- {cert}')
        lines.append('')

    if resume.awards:
        lines.append('AWARDS')
        for award in resume.awards:
            lines.append(f'- {award}')
        lines.append('')

    text_path = output_dir / filename
    text_path.write_text('\n'.join(lines).strip() + '\n', encoding='utf-8')
    return text_path


def render_docx(resume: CanonicalResume, output_dir: Path, filename: str = 'resume.docx') -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.styles['Normal'].font.name = 'Calibri'
    document.styles['Normal'].font.size = Pt(11)

    title = document.add_paragraph(resume.identity.name)
    title.runs[0].bold = True
    title.alignment = 1

    contact = ' | '.join(
        [part for part in [resume.identity.email, resume.identity.phone, resume.identity.location, *resume.identity.links] if _clean_text(part)]
    )
    if contact:
        contact_p = document.add_paragraph(contact)
        contact_p.alignment = 1

    if resume.summary:
        document.add_heading('Summary', level=2)
        document.add_paragraph(_clean_text(resume.summary))

    if resume.education:
        document.add_heading('Education', level=2)
        for entry in resume.education:
            date_text = ' - '.join([part for part in [entry.dates.start, entry.dates.end] if _clean_text(part)])
            header = f'{entry.school}'
            if date_text:
                header = f'{header} ({date_text})'
            document.add_paragraph(header)
            degree = entry.degree
            if entry.major:
                degree = f'{degree} in {entry.major}'
            if degree.strip():
                document.add_paragraph(degree)
            if entry.coursework:
                document.add_paragraph(f'Coursework: {", ".join(entry.coursework)}')

    if resume.experience:
        document.add_heading('Experience', level=2)
        for entry in resume.experience:
            date_text = ' - '.join([part for part in [entry.dates.start, entry.dates.end] if _clean_text(part)])
            document.add_paragraph(f'{entry.title} | {entry.company}' + (f' ({date_text})' if date_text else ''))
            for bullet in entry.bullets:
                document.add_paragraph(bullet, style='List Bullet')

    if resume.projects:
        document.add_heading('Projects', level=2)
        for entry in resume.projects:
            date_text = ''
            if entry.dates:
                date_text = ' - '.join([part for part in [entry.dates.start, entry.dates.end] if _clean_text(part)])
            document.add_paragraph(entry.name + (f' ({date_text})' if date_text else ''))
            if entry.tech:
                document.add_paragraph(f'Tech: {", ".join(entry.tech)}')
            for bullet in entry.bullets:
                document.add_paragraph(bullet, style='List Bullet')

    if resume.skills.categories:
        document.add_heading('Skills', level=2)
        for category, entries in resume.skills.categories.items():
            if _is_generic_skill_category(category):
                document.add_paragraph(', '.join(entries))
            else:
                document.add_paragraph(f'{category}: {", ".join(entries)}')

    if resume.certifications:
        document.add_heading('Certifications', level=2)
        for cert in resume.certifications:
            document.add_paragraph(cert, style='List Bullet')

    if resume.awards:
        document.add_heading('Awards', level=2)
        for award in resume.awards:
            document.add_paragraph(award, style='List Bullet')

    path = output_dir / filename
    document.save(str(path))
    return path


def _escape_pdf_text(value: str) -> str:
    return value.replace('\\', '\\\\').replace('(', '\\(').replace(')', '\\)')


def _build_simple_text_pdf(lines: Sequence[str], path: Path) -> None:
    page_width = 612
    page_height = 792
    margin_left = 42
    margin_top = 760
    line_height = 13
    max_lines_per_page = 52

    pages: List[List[str]] = []
    current: List[str] = []
    for line in lines:
        current.append(line)
        if len(current) >= max_lines_per_page:
            pages.append(current)
            current = []
    if current or not pages:
        pages.append(current or [''])

    objects: List[str] = []
    objects.append('<< /Type /Catalog /Pages 2 0 R >>')

    page_object_ids: List[int] = []
    content_object_ids: List[int] = []
    next_obj_id = 3

    for _ in pages:
        page_object_ids.append(next_obj_id)
        next_obj_id += 1
        content_object_ids.append(next_obj_id)
        next_obj_id += 1

    font_object_id = next_obj_id

    kids = ' '.join(f'{obj_id} 0 R' for obj_id in page_object_ids)
    objects.append(f'<< /Type /Pages /Kids [{kids}] /Count {len(page_object_ids)} >>')

    for idx, page_lines in enumerate(pages):
        page_obj_id = page_object_ids[idx]
        content_obj_id = content_object_ids[idx]

        stream = ['BT', '/F1 10 Tf', f'{margin_left} {margin_top} Td']
        for line_index, raw_line in enumerate(page_lines):
            if line_index > 0:
                stream.append(f'0 -{line_height} Td')
            stream.append(f'({_escape_pdf_text(raw_line)}) Tj')
        stream.append('ET')
        content = '\n'.join(stream) + '\n'

        page_obj = (
            '<< /Type /Page /Parent 2 0 R '
            f'/MediaBox [0 0 {page_width} {page_height}] '
            f'/Resources << /Font << /F1 {font_object_id} 0 R >> >> '
            f'/Contents {content_obj_id} 0 R >>'
        )
        while len(objects) < page_obj_id - 1:
            objects.append('<<>>')
        objects.append(page_obj)

        content_obj = f'<< /Length {len(content.encode("latin-1", errors="ignore"))} >>\nstream\n{content}endstream'
        while len(objects) < content_obj_id - 1:
            objects.append('<<>>')
        objects.append(content_obj)

    while len(objects) < font_object_id - 1:
        objects.append('<<>>')
    objects.append('<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>')

    chunks: List[str] = ['%PDF-1.4\n']
    offsets: List[int] = [0]
    for obj_id, obj_text in enumerate(objects, start=1):
        offsets.append(sum(len(chunk.encode('latin-1', errors='ignore')) for chunk in chunks))
        chunks.append(f'{obj_id} 0 obj\n{obj_text}\nendobj\n')

    xref_offset = sum(len(chunk.encode('latin-1', errors='ignore')) for chunk in chunks)
    chunks.append(f'xref\n0 {len(objects) + 1}\n')
    chunks.append('0000000000 65535 f \n')
    for offset in offsets[1:]:
        chunks.append(f'{offset:010d} 00000 n \n')
    chunks.append(f'trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF\n')
    path.write_bytes(''.join(chunks).encode('latin-1', errors='ignore'))


def render_pdf(
    resume: CanonicalResume,
    output_dir: Path,
    filename: str = 'resume.pdf',
    txt_source_path: Optional[Path] = None,
) -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    txt_path = txt_source_path or render_txt(resume, output_dir)
    lines = [line.rstrip() for line in txt_path.read_text(encoding='utf-8').splitlines()]
    pdf_path = output_dir / filename
    _build_simple_text_pdf(lines, pdf_path)
    return pdf_path


def verify_text_layer(pdf_path: Path, min_chars: int = 80) -> Dict[str, Any]:
    if not pdf_path.exists():
        return {'ok': False, 'extractable_chars': 0, 'reason': 'missing_file'}
    try:
        with pdfplumber.open(pdf_path) as pdf:
            extracted = '\n'.join((page.extract_text() or '') for page in pdf.pages)
    except Exception as exc:
        return {'ok': False, 'extractable_chars': 0, 'reason': f'pdf_parse_error:{exc}'}
    chars = len(_clean_text(extracted))
    return {'ok': chars >= min_chars, 'extractable_chars': chars, 'reason': 'ok' if chars >= min_chars else 'insufficient_text'}


def validate_contact(resume: CanonicalResume) -> Dict[str, Any]:
    identity = resume.identity
    email_ok = bool(re.fullmatch(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}', identity.email or ''))
    phone_digits = re.sub(r'\D', '', identity.phone or '')
    phone_ok = 10 <= len(phone_digits) <= 15

    url_checks: List[Dict[str, Any]] = []
    for link in identity.links:
        parsed_ok = bool(re.match(r'^(https?://)', link.strip()))
        url_checks.append({'url': link, 'valid': parsed_ok})

    return {
        'email_valid': email_ok,
        'phone_valid': phone_ok,
        'urls': url_checks,
        'all_valid': email_ok and phone_ok and all(item['valid'] for item in url_checks),
    }


def detect_sensitive_data(raw_text: str, resume: Optional[CanonicalResume] = None, file_path: Optional[Path] = None) -> Dict[str, Any]:
    findings: List[Dict[str, Any]] = []
    lowered = raw_text.lower()
    if re.search(r'\b(?:dob|date of birth)\b', lowered):
        findings.append(
            _issue(
                ISSUE_SENSITIVE_DATA,
                'high',
                'DOB/date-of-birth text detected.',
                'Remove DOB from resume for ATS and privacy safety.',
            )
        )
    if re.search(r'\b\d{3}-\d{2}-\d{4}\b', raw_text):
        findings.append(
            _issue(
                ISSUE_SENSITIVE_DATA,
                'high',
                'SSN-like pattern detected.',
                'Remove SSN or similar identifiers from resume.',
            )
        )
    if re.search(r'\bpassport\b', lowered):
        findings.append(
            _issue(
                ISSUE_SENSITIVE_DATA,
                'medium',
                'Passport-related sensitive data mentioned.',
                'Remove passport identifiers unless explicitly requested.',
            )
        )

    if file_path and file_path.suffix.lower() == '.docx':
        try:
            doc = Document(str(file_path))
            if len(doc.inline_shapes) > 0:
                findings.append(
                    _issue(
                        ISSUE_SENSITIVE_DATA,
                        'medium',
                        'Embedded image/photo detected in DOCX.',
                        'Remove profile photos for ATS compatibility.',
                        evidence=f'inline_shapes={len(doc.inline_shapes)}',
                    )
                )
        except Exception:
            pass
    if file_path and file_path.suffix.lower() == '.pdf':
        try:
            with pdfplumber.open(file_path) as pdf:
                image_count = sum(len(page.images) for page in pdf.pages)
            if image_count > 0:
                findings.append(
                    _issue(
                        ISSUE_SENSITIVE_DATA,
                        'low',
                        'PDF contains images which may include non-ATS-safe visual elements.',
                        'Prefer text-only ATS-safe export.',
                        evidence=f'image_objects={image_count}',
                    )
                )
        except Exception:
            pass

    if resume:
        contact_check = validate_contact(resume)
        if not contact_check['all_valid']:
            findings.append(
                _issue(
                    ISSUE_SENSITIVE_DATA,
                    'low',
                    'Contact block has invalid fields.',
                    'Fix malformed email/phone/URL values in identity.',
                )
            )

    return {'findings': findings, 'count': len(findings)}


def _parse_date_value(value: Optional[str]) -> Dict[str, Any]:
    text = _clean_text(value or '')
    if not text:
        return {'normalized': None, 'year': None, 'month': None, 'is_present': False, 'unknown_month': True}
    lowered = text.lower()
    if lowered == 'present':
        now = datetime.now(timezone.utc)
        return {
            'normalized': f'{now.year:04d}-{now.month:02d}',
            'year': now.year,
            'month': now.month,
            'is_present': True,
            'unknown_month': False,
        }

    month_match = re.search(
        r'\b(jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|jun(?:e)?|jul(?:y)?|aug(?:ust)?|'
        r'sep(?:t|tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)\b[\s,/-]*(\d{4})',
        lowered,
    )
    if month_match:
        month = MONTHS[month_match.group(1)[:3]]
        year = int(month_match.group(2))
        return {
            'normalized': f'{year:04d}-{month:02d}',
            'year': year,
            'month': month,
            'is_present': False,
            'unknown_month': False,
        }

    year_match = re.search(r'\b(19|20)\d{2}\b', lowered)
    if year_match:
        year = int(year_match.group(0))
        return {
            'normalized': f'{year:04d}-??',
            'year': year,
            'month': None,
            'is_present': False,
            'unknown_month': True,
        }
    return {'normalized': None, 'year': None, 'month': None, 'is_present': False, 'unknown_month': True}


def normalize_dates(resume: CanonicalResume) -> Dict[str, Any]:
    def normalize_range(label: str, dates: Optional[DateRange]) -> Dict[str, Any]:
        start_info = _parse_date_value(dates.start if dates else None)
        end_info = _parse_date_value(dates.end if dates else None)
        return {
            'label': label,
            'start': start_info['normalized'],
            'end': end_info['normalized'],
            'start_unknown_month': start_info['unknown_month'],
            'end_unknown_month': end_info['unknown_month'],
            'start_meta': start_info,
            'end_meta': end_info,
        }

    experience = [normalize_range(f'{entry.title}@{entry.company}', entry.dates) for entry in resume.experience]
    education = [normalize_range(entry.school, entry.dates) for entry in resume.education]
    projects = [normalize_range(entry.name, entry.dates) for entry in resume.projects if entry.dates]
    return {'experience': experience, 'education': education, 'projects': projects}


def _month_bounds(meta: Dict[str, Any], is_start: bool) -> Tuple[Optional[int], Optional[int]]:
    year = meta.get('year')
    month = meta.get('month')
    if year is None:
        return None, None
    if month is not None:
        value = (year * 12) + month
        return value, value
    if is_start:
        return (year * 12) + 1, (year * 12) + 12
    return (year * 12) + 1, (year * 12) + 12


def detect_overlaps(normalized_dates: Dict[str, Any]) -> Dict[str, Any]:
    overlaps: List[Dict[str, Any]] = []
    experience = normalized_dates.get('experience', [])
    intervals: List[Tuple[str, int, int]] = []
    for row in experience:
        start_min, start_max = _month_bounds(row.get('start_meta', {}), is_start=True)
        end_min, end_max = _month_bounds(row.get('end_meta', {}), is_start=False)
        if start_min is None or end_max is None:
            continue
        intervals.append((row.get('label', ''), start_min, end_max))

    for idx in range(len(intervals)):
        left_label, left_start, left_end = intervals[idx]
        for jdx in range(idx + 1, len(intervals)):
            right_label, right_start, right_end = intervals[jdx]
            if left_start <= right_end and right_start <= left_end:
                overlaps.append({'left': left_label, 'right': right_label})

    return {'overlaps': overlaps, 'count': len(overlaps)}


def compute_durations(normalized_dates: Dict[str, Any]) -> Dict[str, Any]:
    durations: List[Dict[str, Any]] = []
    total_min = 0
    total_max = 0
    for row in normalized_dates.get('experience', []):
        start_meta = row.get('start_meta', {})
        end_meta = row.get('end_meta', {})
        start_min, start_max = _month_bounds(start_meta, is_start=True)
        end_min, end_max = _month_bounds(end_meta, is_start=False)
        if start_max is None or end_min is None:
            continue
        min_months = max(0, end_min - start_max + 1)
        max_months = max(0, (end_max or end_min) - (start_min or start_max) + 1)
        durations.append({'label': row.get('label', ''), 'months_min': min_months, 'months_max': max_months})
        total_min += min_months
        total_max += max_months

    return {
        'experience_durations': durations,
        'total_experience_months_min': total_min,
        'total_experience_months_max': total_max,
        'total_experience_years_min': round(total_min / 12.0, 2),
        'total_experience_years_max': round(total_max / 12.0, 2),
    }


def maintain_skill_alias_graph(additional_aliases: Optional[Dict[str, List[str]]] = None) -> Dict[str, List[str]]:
    return _canonical_skill_graph(additional_aliases)


def extract_skills(
    text: str,
    alias_graph: Optional[Dict[str, List[str]]] = None,
    soft_alias_graph: Optional[Dict[str, List[str]]] = None,
) -> Dict[str, Any]:
    hard_graph = _canonical_skill_graph(alias_graph)
    soft_graph = _canonical_soft_skill_graph(soft_alias_graph)
    raw_skills = _extract_alias_matches(text, hard_graph)
    raw_soft_skills = _extract_alias_matches(text, soft_graph)

    title_candidates = re.findall(r'[A-Z][A-Za-z/&\-\s]{2,60}', text)
    titles: List[str] = []
    for candidate in title_candidates:
        tokens = tokenize(candidate)
        if any(token in TITLE_TOKENS for token in tokens):
            cleaned = _clean_text(candidate)
            if cleaned and cleaned not in titles:
                titles.append(cleaned)

    return {'skills': raw_skills, 'soft_skills': raw_soft_skills, 'titles': titles}


def map_to_canonical_skills(raw_skills: Sequence[Dict[str, Any]], alias_graph: Optional[Dict[str, List[str]]] = None) -> List[Dict[str, Any]]:
    _ = _canonical_skill_graph(alias_graph)
    grouped: Dict[str, Dict[str, Any]] = {}
    for row in raw_skills:
        canonical_id = _normalize_canonical_id(row.get('canonical_id', ''))
        if not canonical_id:
            continue
        bucket = grouped.setdefault(canonical_id, {'canonical_id': canonical_id, 'evidence_spans': [], 'matched_aliases': []})
        span = row.get('evidence_span')
        if isinstance(span, list) and len(span) == 2:
            bucket['evidence_spans'].append(span)
        alias = _clean_text(str(row.get('matched_text', '')))
        if alias and alias not in bucket['matched_aliases']:
            bucket['matched_aliases'].append(alias)
    return [grouped[key] for key in sorted(grouped.keys())]


def parse_job_description(jd_text: str, alias_graph: Optional[Dict[str, List[str]]] = None) -> Dict[str, Any]:
    lines = [line.strip() for line in jd_text.splitlines() if line.strip()]
    required_lines: List[str] = []
    preferred_lines: List[str] = []
    other_lines: List[str] = []
    mode = 'other'
    required_markers = [
        'required',
        'must have',
        'minimum qualification',
        'minimum requirement',
        'basic qualification',
    ]
    preferred_markers = [
        'preferred',
        'nice to have',
        'bonus',
        'preferred qualification',
    ]

    def _inline_suffix(line: str, marker: str) -> str:
        lowered_line = line.lower()
        idx = lowered_line.find(marker)
        if idx < 0:
            return ''
        suffix = _clean_text(line[idx + len(marker):].lstrip(': -'))
        if normalize_token(suffix) in {'qualification', 'qualifications', 'requirement', 'requirements', 'skills'}:
            return ''
        return suffix

    for line in lines:
        lowered = line.lower()
        required_marker = next((marker for marker in required_markers if marker in lowered), None)
        if required_marker:
            mode = 'required'
            suffix = _inline_suffix(line, required_marker)
            if suffix:
                required_lines.append(suffix)
            continue
        preferred_marker = next((marker for marker in preferred_markers if marker in lowered), None)
        if preferred_marker:
            mode = 'preferred'
            suffix = _inline_suffix(line, preferred_marker)
            if suffix:
                preferred_lines.append(suffix)
            continue

        if mode == 'required':
            required_lines.append(line)
        elif mode == 'preferred':
            preferred_lines.append(line)
        else:
            other_lines.append(line)

    required_text = ' '.join(required_lines)
    preferred_text = ' '.join(preferred_lines)
    all_text = ' '.join(lines)

    required_raw = extract_skills(required_text or all_text, alias_graph=alias_graph)
    preferred_raw = extract_skills(preferred_text, alias_graph=alias_graph)

    required_skills = map_to_canonical_skills(required_raw['skills'], alias_graph=alias_graph)
    preferred_skills = map_to_canonical_skills(preferred_raw['skills'], alias_graph=alias_graph)
    required_soft_skills = map_to_canonical_skills(required_raw.get('soft_skills', []))
    preferred_soft_skills = map_to_canonical_skills(preferred_raw.get('soft_skills', []))

    years = [int(value) for value in re.findall(r'\b(\d{1,2})\+?\s+years?\b', all_text.lower())]
    years_required = max(years) if years else None

    degrees: List[str] = []
    lowered_all = all_text.lower()
    for degree_key, aliases in DEGREE_TOKENS.items():
        if any(alias in lowered_all for alias in aliases):
            degrees.append(degree_key)

    certs = re.findall(
        r'\b(?:aws|gcp|azure|pmp|cissp|security\+|cka|ckad|scrum master|comptia)\b',
        all_text,
        flags=re.IGNORECASE,
    )
    certifications = unique_preserve_order([_clean_text(cert).upper() for cert in certs])

    location_match = re.search(r'\b(?:location|based in|onsite|hybrid|remote)\b[:\-\s]*([A-Za-z ,]+)', all_text, flags=re.IGNORECASE)
    location = _clean_text(location_match.group(1)) if location_match else None

    work_auth = None
    if re.search(r'work authorization|required to work|us citizen|visa sponsorship', lowered_all):
        work_auth = 'restricted_or_specified'

    role_tokens = [token for token in tokenize(all_text) if token in TITLE_TOKENS]
    normalized_title = None
    if 'engineer' in role_tokens and 'ml' in tokenize(all_text.lower()):
        normalized_title = 'machine_learning_engineer'
    elif role_tokens:
        normalized_title = '_'.join(role_tokens[:2])

    return {
        'required': {
            'skills': required_skills,
            'soft_skills': required_soft_skills,
            'lines': required_lines,
        },
        'preferred': {
            'skills': preferred_skills,
            'soft_skills': preferred_soft_skills,
            'lines': preferred_lines,
        },
        'other_lines': other_lines,
        'years_required': years_required,
        'degree_constraints': sorted(set(degrees)),
        'cert_constraints': certifications,
        'location_constraint': location,
        'work_auth_constraint': work_auth,
        'normalized_title': normalized_title,
    }


def build_requirement_graph(parsed_job: Dict[str, Any], alias_graph: Optional[Dict[str, List[str]]] = None) -> Dict[str, Any]:
    _ = _canonical_skill_graph(alias_graph)
    nodes: List[Dict[str, Any]] = [{'id': 'job', 'type': 'job', 'label': 'Target Job'}]
    edges: List[Dict[str, Any]] = []
    node_ids: Set[str] = {'job'}

    def add_skill_nodes(skills: Sequence[Dict[str, Any]], required: bool, soft: bool) -> None:
        relation = 'requires' if required else 'prefers'
        node_type = 'soft_skill' if soft else 'hard_skill'
        for skill in skills:
            canonical_id = _clean_text(str(skill.get('canonical_id', '')))
            if not canonical_id:
                continue
            node_id = f"{node_type}:{canonical_id}"
            if node_id not in node_ids:
                nodes.append({'id': node_id, 'type': node_type, 'label': canonical_id, 'required': required})
                node_ids.add(node_id)
            edges.append({'source': 'job', 'target': node_id, 'relation': relation})

    add_skill_nodes(parsed_job.get('required', {}).get('skills', []), required=True, soft=False)
    add_skill_nodes(parsed_job.get('required', {}).get('soft_skills', []), required=True, soft=True)
    add_skill_nodes(parsed_job.get('preferred', {}).get('skills', []), required=False, soft=False)
    add_skill_nodes(parsed_job.get('preferred', {}).get('soft_skills', []), required=False, soft=True)

    constraints = {
        'years_required': parsed_job.get('years_required'),
        'degree_constraints': parsed_job.get('degree_constraints', []),
        'cert_constraints': parsed_job.get('cert_constraints', []),
        'location_constraint': parsed_job.get('location_constraint'),
        'work_auth_constraint': parsed_job.get('work_auth_constraint'),
        'normalized_title': parsed_job.get('normalized_title'),
    }
    return {'nodes': nodes, 'edges': edges, 'constraints': constraints}


def _hash_embedding(text: str, dims: int = 128) -> List[float]:
    vector = [0.0] * dims
    for token in tokenize(text):
        if not token:
            continue
        digest_int = int(hashlib.sha256(token.encode('utf-8')).hexdigest()[:16], 16)
        index = digest_int % dims
        sign = -1.0 if (digest_int // dims) % 2 else 1.0
        vector[index] += sign
    return vector


def _cosine_similarity(left: Sequence[float], right: Sequence[float]) -> float:
    numerator = sum(a * b for a, b in zip(left, right))
    left_norm = math.sqrt(sum(a * a for a in left))
    right_norm = math.sqrt(sum(b * b for b in right))
    if left_norm == 0 or right_norm == 0:
        return 0.0
    return numerator / (left_norm * right_norm)


def _resume_text(resume: CanonicalResume) -> str:
    chunks = [
        resume.identity.name,
        resume.identity.location,
        resume.summary or '',
    ]
    for entry in resume.experience:
        chunks.extend([entry.company, entry.title, *entry.bullets])
    for entry in resume.projects:
        chunks.extend([entry.name, *entry.tech, *entry.bullets])
    for entry in resume.education:
        chunks.extend([entry.school, entry.degree, entry.major, *entry.coursework])
    for values in resume.skills.categories.values():
        chunks.extend(values)
    chunks.extend(resume.certifications)
    chunks.extend(resume.awards or [])
    return ' '.join(chunks)


def _resume_section_texts(resume: CanonicalResume) -> Dict[str, str]:
    sections: Dict[str, List[str]] = {
        'identity': [resume.identity.name, resume.identity.email, resume.identity.phone, resume.identity.location, *resume.identity.links],
        'summary': [resume.summary or ''],
        'education': [],
        'experience': [],
        'projects': [],
        'skills': [],
        'certifications': list(resume.certifications),
    }
    for entry in resume.education:
        sections['education'].extend([entry.school, entry.degree, entry.major, *entry.minors, *entry.coursework])
    for entry in resume.experience:
        sections['experience'].extend([entry.company, entry.title, entry.location, *entry.bullets])
    for entry in resume.projects:
        sections['projects'].extend([entry.name, *(entry.tech or []), *(entry.bullets or [])])
    for category, entries in resume.skills.categories.items():
        sections['skills'].append(category)
        sections['skills'].extend(entries)
    return {name: _clean_text(' '.join(values)).lower() for name, values in sections.items()}


def _resume_skill_evidence(
    resume: CanonicalResume,
    alias_graph: Optional[Dict[str, List[str]]] = None,
    *,
    soft: bool = False,
) -> Dict[str, Dict[str, Any]]:
    graph = _canonical_soft_skill_graph(alias_graph) if soft else _canonical_skill_graph(alias_graph)
    section_texts = _resume_section_texts(resume)
    evidence: Dict[str, Dict[str, Any]] = {}
    for canonical_id, aliases in sorted(graph.items()):
        matched_sections: List[str] = []
        best_weight = 0.0
        for section_name, section_text in section_texts.items():
            if not section_text:
                continue
            matched = False
            for alias in aliases:
                if not alias:
                    continue
                if re.search(rf'(?<![A-Za-z0-9]){re.escape(alias)}(?![A-Za-z0-9])', section_text):
                    matched = True
                    break
            if matched:
                matched_sections.append(section_name)
                best_weight = max(best_weight, SECTION_EVIDENCE_WEIGHTS.get(section_name, 0.4))
        if best_weight > 0.0:
            evidence[canonical_id] = {
                'weight': round(best_weight, 3),
                'sections': matched_sections,
            }
    return evidence


def _resume_skill_ids(resume: CanonicalResume, alias_graph: Optional[Dict[str, List[str]]] = None) -> Set[str]:
    return set(_resume_skill_evidence(resume, alias_graph=alias_graph, soft=False).keys())


def _resume_soft_skill_ids(resume: CanonicalResume, alias_graph: Optional[Dict[str, List[str]]] = None) -> Set[str]:
    return set(_resume_skill_evidence(resume, alias_graph=alias_graph, soft=True).keys())


def _resume_years_estimate(resume: CanonicalResume) -> float:
    normalized = normalize_dates(resume)
    durations = compute_durations(normalized)
    return durations['total_experience_years_min']


def _degree_satisfied(resume: CanonicalResume, degree_constraints: Sequence[str]) -> bool:
    if not degree_constraints:
        return True
    degrees_text = ' '.join(entry.degree for entry in resume.education).lower()
    for constraint in degree_constraints:
        aliases = DEGREE_TOKENS.get(constraint, {constraint})
        if any(alias in degrees_text for alias in aliases):
            return True
    return False


def compute_match_score(
    resume: CanonicalResume,
    jd_text: str,
    alias_graph: Optional[Dict[str, List[str]]] = None,
) -> Dict[str, Any]:
    parsed_jd = parse_job_description(jd_text, alias_graph=alias_graph)
    requirement_graph = build_requirement_graph(parsed_jd, alias_graph=alias_graph)

    required_hard_skills = {row['canonical_id'] for row in parsed_jd.get('required', {}).get('skills', [])}
    preferred_hard_skills = {row['canonical_id'] for row in parsed_jd.get('preferred', {}).get('skills', [])}
    required_soft_skills = {row['canonical_id'] for row in parsed_jd.get('required', {}).get('soft_skills', [])}
    preferred_soft_skills = {row['canonical_id'] for row in parsed_jd.get('preferred', {}).get('soft_skills', [])}

    resume_hard_evidence = _resume_skill_evidence(resume, alias_graph=alias_graph, soft=False)
    resume_soft_evidence = _resume_skill_evidence(resume, soft=True)
    resume_hard_skills = set(resume_hard_evidence.keys())
    resume_soft_skills = set(resume_soft_evidence.keys())

    overlap_required_hard = sorted(required_hard_skills & resume_hard_skills)
    overlap_preferred_hard = sorted(preferred_hard_skills & resume_hard_skills)
    missing_required_hard = sorted(required_hard_skills - resume_hard_skills)

    overlap_required_soft = sorted(required_soft_skills & resume_soft_skills)
    overlap_preferred_soft = sorted(preferred_soft_skills & resume_soft_skills)
    missing_required_soft = sorted(required_soft_skills - resume_soft_skills)

    hard_required_coverage = len(overlap_required_hard) / max(1, len(required_hard_skills)) if required_hard_skills else 1.0
    hard_preferred_coverage = len(overlap_preferred_hard) / max(1, len(preferred_hard_skills)) if preferred_hard_skills else 1.0
    hard_required_weighted_coverage = (
        sum(float(resume_hard_evidence.get(skill, {}).get('weight', 0.0)) for skill in required_hard_skills) / max(1, len(required_hard_skills))
        if required_hard_skills
        else 1.0
    )
    hard_preferred_weighted_coverage = (
        sum(float(resume_hard_evidence.get(skill, {}).get('weight', 0.0)) for skill in preferred_hard_skills) / max(1, len(preferred_hard_skills))
        if preferred_hard_skills
        else 1.0
    )
    hard_skill_score = min(100.0, (hard_required_weighted_coverage * 85.0) + (hard_preferred_weighted_coverage * 15.0))

    soft_required_coverage = len(overlap_required_soft) / max(1, len(required_soft_skills)) if required_soft_skills else 1.0
    soft_preferred_coverage = len(overlap_preferred_soft) / max(1, len(preferred_soft_skills)) if preferred_soft_skills else 1.0
    soft_required_weighted_coverage = (
        sum(float(resume_soft_evidence.get(skill, {}).get('weight', 0.0)) for skill in required_soft_skills) / max(1, len(required_soft_skills))
        if required_soft_skills
        else 1.0
    )
    soft_preferred_weighted_coverage = (
        sum(float(resume_soft_evidence.get(skill, {}).get('weight', 0.0)) for skill in preferred_soft_skills) / max(1, len(preferred_soft_skills))
        if preferred_soft_skills
        else 1.0
    )
    soft_skill_score = min(100.0, (soft_required_weighted_coverage * 85.0) + (soft_preferred_weighted_coverage * 15.0))

    hard_weight = (len(required_hard_skills) * 2) + len(preferred_hard_skills)
    soft_weight = (len(required_soft_skills) * 2) + len(preferred_soft_skills)
    if hard_weight == 0 and soft_weight == 0:
        lexical_score = 50.0
    elif soft_weight == 0:
        lexical_score = hard_skill_score
    elif hard_weight == 0:
        lexical_score = soft_skill_score
    else:
        lexical_score = ((hard_skill_score * hard_weight) + (soft_skill_score * soft_weight)) / (hard_weight + soft_weight)
    lexical_score = min(100.0, max(0.0, lexical_score))

    semantic_similarity = _cosine_similarity(_hash_embedding(_resume_text(resume)), _hash_embedding(jd_text))
    semantic_score = max(0.0, min(100.0, ((semantic_similarity + 1.0) / 2.0) * 100.0))

    years_required = parsed_jd.get('years_required')
    resume_years = _resume_years_estimate(resume)
    if years_required is None:
        years_score = 100.0
    elif resume_years >= years_required:
        years_score = 100.0
    elif years_required > 0:
        years_score = max(0.0, (resume_years / years_required) * 100.0)
    else:
        years_score = 100.0

    degree_score = 100.0 if _degree_satisfied(resume, parsed_jd.get('degree_constraints', [])) else 30.0
    location_constraint = _clean_text(parsed_jd.get('location_constraint') or '')
    if not location_constraint:
        location_score = 100.0
    else:
        location_score = 100.0 if location_constraint.lower() in resume.identity.location.lower() else 60.0
    work_auth_score = 70.0 if parsed_jd.get('work_auth_constraint') else 100.0

    structured_score = (years_score * 0.45) + (degree_score * 0.25) + (location_score * 0.2) + (work_auth_score * 0.1)
    overall = (lexical_score * 0.4) + (semantic_score * 0.35) + (structured_score * 0.25)

    must_have_gate_applied = False
    must_have_gate_cap = 100.0
    if len(required_hard_skills) >= MUST_HAVE_MIN_REQUIRED_HARD_SKILLS and hard_required_coverage < MUST_HAVE_REQUIRED_HARD_COVERAGE:
        must_have_gate_applied = True
        must_have_gate_cap = max(25.0, 35.0 + (hard_required_coverage * 65.0))
        overall = min(overall, must_have_gate_cap)

    top_drivers: List[Dict[str, Any]] = []
    for skill in overlap_required_hard[:5]:
        top_drivers.append({'type': 'required_hard_skill_match', 'value': skill, 'impact': 8.0})
    for skill in overlap_required_soft[:3]:
        top_drivers.append({'type': 'required_soft_skill_match', 'value': skill, 'impact': 5.0})
    if years_required is not None:
        top_drivers.append({'type': 'years_experience', 'value': {'resume_years': resume_years, 'required_years': years_required}, 'impact': 6.0})
    top_drivers.append({'type': 'degree_match', 'value': degree_score >= 100.0, 'impact': 4.0})

    ranked_gaps: List[Dict[str, Any]] = []
    for idx, skill in enumerate(missing_required_hard):
        ranked_gaps.append({'rank': idx + 1, 'type': 'missing_required_hard_skill', 'value': skill, 'severity': 'high'})
    for skill in missing_required_soft:
        ranked_gaps.append(
            {
                'rank': len(ranked_gaps) + 1,
                'type': 'missing_required_soft_skill',
                'value': skill,
                'severity': 'medium',
            }
        )
    if years_required is not None and resume_years < years_required:
        ranked_gaps.append(
            {
                'rank': len(ranked_gaps) + 1,
                'type': 'years_shortfall',
                'value': {'resume_years': resume_years, 'required_years': years_required},
                'severity': 'medium',
            }
        )
    if must_have_gate_applied:
        ranked_gaps.append(
            {
                'rank': len(ranked_gaps) + 1,
                'type': 'must_have_coverage_shortfall',
                'value': {
                    'required_hard_coverage': round(hard_required_coverage, 3),
                    'threshold': MUST_HAVE_REQUIRED_HARD_COVERAGE,
                    'score_cap': round(must_have_gate_cap, 2),
                },
                'severity': 'high',
            }
        )

    subscores = {
        'lexical_similarity': round(lexical_score, 2),
        'semantic_similarity': round(semantic_score, 2),
        'structured_constraints': round(structured_score, 2),
        'hard_skill_alignment': round(hard_skill_score, 2),
        'soft_skill_alignment': round(soft_skill_score, 2),
    }

    return {
        'overall_score': round(max(0.0, min(100.0, overall)), 2),
        'subscores': subscores,
        'top_drivers': top_drivers,
        'ranked_gaps': ranked_gaps,
        'metadata': {
            'required_skills': sorted(required_hard_skills),
            'matched_required_skills': overlap_required_hard,
            'matched_preferred_skills': overlap_preferred_hard,
            'required_hard_skills': sorted(required_hard_skills),
            'preferred_hard_skills': sorted(preferred_hard_skills),
            'required_soft_skills': sorted(required_soft_skills),
            'preferred_soft_skills': sorted(preferred_soft_skills),
            'matched_required_hard_skills': overlap_required_hard,
            'matched_preferred_hard_skills': overlap_preferred_hard,
            'matched_required_soft_skills': overlap_required_soft,
            'matched_preferred_soft_skills': overlap_preferred_soft,
            'required_hard_coverage': round(hard_required_coverage, 3),
            'required_hard_weighted_coverage': round(hard_required_weighted_coverage, 3),
            'required_soft_coverage': round(soft_required_coverage, 3),
            'required_soft_weighted_coverage': round(soft_required_weighted_coverage, 3),
            'must_have_gate_applied': must_have_gate_applied,
            'must_have_gate_cap': round(must_have_gate_cap, 2),
            'hard_skill_evidence': resume_hard_evidence,
            'soft_skill_evidence': resume_soft_evidence,
            'resume_hard_skill_ids': sorted(resume_hard_skills),
            'resume_soft_skill_ids': sorted(resume_soft_skills),
            'resume_skill_ids': sorted(resume_hard_skills),
            'requirement_graph': requirement_graph,
        },
    }


def generate_match_explanations(match_score: Dict[str, Any]) -> Dict[str, Any]:
    subscores = match_score.get('subscores', {})
    drivers = match_score.get('top_drivers', [])
    gaps = match_score.get('ranked_gaps', [])
    return {
        'overall_score': match_score.get('overall_score', 0.0),
        'subscores': subscores,
        'drivers': drivers,
        'gaps': gaps,
        'explainable': True,
    }


def _summary_from_existing_terms(resume: CanonicalResume, jd_text: str, alias_graph: Optional[Dict[str, List[str]]] = None) -> str:
    parsed = parse_job_description(jd_text, alias_graph=alias_graph)
    required = [row['canonical_id'] for row in parsed['required']['skills']]
    resume_skills = _resume_skill_ids(resume, alias_graph=alias_graph)
    aligned = [skill for skill in required if skill in resume_skills]
    fallback = [skill for skill in sorted(resume_skills) if skill in DEFAULT_SKILL_ALIAS_GRAPH][:3]
    use_terms = aligned[:3] if aligned else fallback[:3]
    readable = [term.replace('_', ' ') for term in use_terms]
    if not readable:
        return resume.summary or 'Engineer with hands-on delivery experience across selected resume projects and work history.'
    if len(readable) == 1:
        terms_text = readable[0]
    elif len(readable) == 2:
        terms_text = f'{readable[0]} and {readable[1]}'
    else:
        terms_text = f"{', '.join(readable[:-1])}, and {readable[-1]}"
    return f'Engineer with hands-on experience in {terms_text} across selected work and project history.'


def generate_patches(
    resume: CanonicalResume,
    jd_text: str,
    alias_graph: Optional[Dict[str, List[str]]] = None,
) -> Dict[str, Any]:
    match_score = compute_match_score(resume, jd_text, alias_graph=alias_graph)
    metadata = match_score.get('metadata', {})
    required_hard_skills = set(metadata.get('required_hard_skills') or metadata.get('required_skills') or [])
    required_soft_skills = set(metadata.get('required_soft_skills') or [])

    def item_skill_overlap(text: str) -> int:
        tokens = set(tokenize(text))
        return len(tokens & required_hard_skills)

    patches: List[Dict[str, Any]] = []

    if resume.experience and required_hard_skills:
        ranked_experience = sorted(
            resume.experience,
            key=lambda entry: (-item_skill_overlap(' '.join([entry.title, entry.company, *entry.bullets])), entry.title.lower()),
        )
        new_order = [entry.title for entry in ranked_experience]
        current_order = [entry.title for entry in resume.experience]
        if new_order != current_order:
            patches.append(
                {
                    'op': 'reorder_experience',
                    'path': '/experience',
                    'value': new_order,
                    'status': 'GROUNDED',
                    'grounded': True,
                    'requires_user_confirmation': False,
                    'reason': 'Prioritize experience entries with stronger required-skill overlap.',
                    'evidence': sorted(required_hard_skills),
                }
            )

    if resume.projects and required_hard_skills:
        ranked_projects = sorted(
            resume.projects,
            key=lambda entry: (-item_skill_overlap(' '.join([entry.name, *entry.tech, *entry.bullets])), entry.name.lower()),
        )
        new_order = [entry.name for entry in ranked_projects]
        current_order = [entry.name for entry in resume.projects]
        if new_order != current_order:
            patches.append(
                {
                    'op': 'reorder_projects',
                    'path': '/projects',
                    'value': new_order,
                    'status': 'GROUNDED',
                    'grounded': True,
                    'requires_user_confirmation': False,
                    'reason': 'Prioritize projects aligned to required skills.',
                    'evidence': sorted(required_hard_skills),
                }
            )

    proposed_summary = _summary_from_existing_terms(resume, jd_text, alias_graph=alias_graph)
    if _clean_text(proposed_summary) and _clean_text(proposed_summary) != _clean_text(resume.summary or ''):
        patches.append(
            {
                'op': 'set_summary',
                'path': '/summary',
                'value': proposed_summary,
                'status': 'GROUNDED',
                'grounded': True,
                'requires_user_confirmation': False,
                'reason': 'Summary generated strictly from existing resume evidence and required skill overlap.',
                'evidence': sorted(required_hard_skills | required_soft_skills),
            }
        )

    for gap in match_score['ranked_gaps'][:5]:
        gap_type = _clean_text(str(gap.get('type', '')))
        if gap_type not in {'missing_required_hard_skill', 'missing_required_soft_skill', 'missing_required_skill'}:
            continue
        requirement_kind = 'soft_skill' if gap_type == 'missing_required_soft_skill' else 'hard_skill'
        reason = (
            'Required soft skill not explicitly evidenced in canonical resume; cannot be added automatically.'
            if requirement_kind == 'soft_skill'
            else 'Required skill not evidenced in canonical resume; cannot be added automatically.'
        )
        patches.append(
            {
                'op': 'flag_missing_requirement',
                'path': '/skills',
                'value': gap['value'],
                'requirement_kind': requirement_kind,
                'status': 'REQUIRES_USER_CONFIRMATION',
                'grounded': False,
                'requires_user_confirmation': True,
                'reason': reason,
                'evidence': [],
            }
        )

    return {
        'patches': patches,
        'match_score': match_score,
        'explanations': generate_match_explanations(match_score),
    }


def apply_patches(
    resume: CanonicalResume,
    patches: Sequence[Dict[str, Any]],
    allow_requires_confirmation: bool = False,
) -> Dict[str, Any]:
    updated = CanonicalResume.model_validate(deepcopy(resume.model_dump()))
    applied: List[Dict[str, Any]] = []
    skipped: List[Dict[str, Any]] = []

    for patch in patches:
        status = str(patch.get('status', '')).strip().upper()
        requires_confirmation = bool(patch.get('requires_user_confirmation')) or status == 'REQUIRES_USER_CONFIRMATION'
        if requires_confirmation and not allow_requires_confirmation:
            skipped.append({'patch': patch, 'reason': 'requires_user_confirmation'})
            continue

        op = patch.get('op')
        value = patch.get('value')
        if op == 'set_summary' and isinstance(value, str):
            updated.summary = value
            applied.append(patch)
            continue

        if op == 'reorder_experience' and isinstance(value, list):
            order = [str(item) for item in value]
            mapping = {entry.title: entry for entry in updated.experience}
            if set(order) == set(mapping.keys()):
                updated.experience = [mapping[key] for key in order]
                applied.append(patch)
            else:
                skipped.append({'patch': patch, 'reason': 'invalid_experience_order'})
            continue

        if op == 'reorder_projects' and isinstance(value, list):
            order = [str(item) for item in value]
            mapping = {entry.name: entry for entry in updated.projects}
            if set(order) == set(mapping.keys()):
                updated.projects = [mapping[key] for key in order]
                applied.append(patch)
            else:
                skipped.append({'patch': patch, 'reason': 'invalid_project_order'})
            continue

        if op == 'flag_missing_requirement':
            skipped.append({'patch': patch, 'reason': 'requires_user_confirmation'})
            continue

        skipped.append({'patch': patch, 'reason': 'unsupported_operation'})

    return {
        'resume': updated,
        'applied': applied,
        'skipped': skipped,
    }


def compare_versions(left: Dict[str, Any], right: Dict[str, Any]) -> Dict[str, Any]:
    diffs: List[Dict[str, Any]] = []

    def walk(path: str, a: Any, b: Any) -> None:
        if type(a) != type(b):
            diffs.append({'path': path, 'left': a, 'right': b})
            return
        if isinstance(a, dict):
            keys = sorted(set(a.keys()) | set(b.keys()))
            for key in keys:
                walk(f'{path}/{key}', a.get(key), b.get(key))
            return
        if isinstance(a, list):
            if a != b:
                diffs.append({'path': path, 'left': a, 'right': b})
            return
        if a != b:
            diffs.append({'path': path, 'left': a, 'right': b})

    walk('', left, right)
    left_score = float(left.get('match_score', 0.0) or 0.0)
    right_score = float(right.get('match_score', 0.0) or 0.0)
    return {
        'diffs': diffs,
        'diff_count': len(diffs),
        'score_delta': round(right_score - left_score, 2),
    }


def version_resume(
    resume: CanonicalResume,
    *,
    data_dir: Path,
    job_id: str,
    match_score: Optional[float] = None,
    metadata: Optional[Dict[str, Any]] = None,
) -> Dict[str, Any]:
    timestamp = datetime.now(timezone.utc).strftime('%Y%m%d-%H%M%S-%f')
    version_id = f'{timestamp}-{normalize_token(job_id) or "job"}'
    root = data_dir / 'versions' / normalize_token(job_id or 'job')
    root.mkdir(parents=True, exist_ok=True)
    payload = {
        'version_id': version_id,
        'job_id': job_id,
        'created_at': datetime.now(timezone.utc).isoformat(),
        'match_score': match_score,
        'resume': resume.model_dump(mode='json'),
        'metadata': metadata or {},
    }
    save_json(root / f'{version_id}.json', payload)
    return payload


def build_canonical(parse_mirror_result: Dict[str, Any]) -> CanonicalResume:
    canonical_payload = parse_mirror_result.get('canonical', parse_mirror_result)
    return CanonicalResume.model_validate(canonical_payload)


def render_outputs(resume: CanonicalResume, output_dir: Path, filename_prefix: str = '') -> Dict[str, Any]:
    safe_prefix = re.sub(r'[^A-Za-z0-9_\-]', '', filename_prefix or '')
    txt_path = render_txt(resume, output_dir, filename=f'{safe_prefix}resume.txt')
    docx_path = render_docx(resume, output_dir, filename=f'{safe_prefix}resume.docx')
    pdf_path = render_pdf(resume, output_dir, filename=f'{safe_prefix}resume.pdf', txt_source_path=txt_path)
    verification = verify_text_layer(pdf_path)
    return {
        'txt_path': str(txt_path),
        'docx_path': str(docx_path),
        'pdf_path': str(pdf_path),
        'pdf_text_layer': verification,
    }


def export_bundle(output_dir: Path, bundle_path: Optional[Path] = None) -> Path:
    if not output_dir.exists():
        raise FileNotFoundError(f'Output directory not found: {output_dir}')
    target = bundle_path or (output_dir / 'bundle.zip')
    target.parent.mkdir(parents=True, exist_ok=True)
    target_resolved = target.resolve()
    with zipfile.ZipFile(target, mode='w', compression=zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(output_dir.rglob('*')):
            if path.is_dir():
                continue
            if path.resolve() == target_resolved:
                continue
            archive.write(path, arcname=path.relative_to(output_dir))
    return target


def upload_resume(
    *,
    file_path: Path,
    enable_ocr: bool,
    llm: Optional[LLMService],
) -> Dict[str, Any]:
    raw_text, extraction_warnings = extract_resume_text(file_path, enable_ocr=enable_ocr)
    mirror = parse_mirror(raw_text, llm=llm)
    canonical = build_canonical(mirror)
    lint = lint_resume(file_path)
    contact = validate_contact(canonical)
    sensitive = detect_sensitive_data(raw_text, resume=canonical, file_path=file_path)
    timeline = normalize_dates(canonical)
    overlaps = detect_overlaps(timeline)
    durations = compute_durations(timeline)
    return {
        'raw_text': raw_text,
        'extraction_warnings': extraction_warnings,
        'parse_mirror': mirror,
        'canonical': canonical.model_dump(mode='json'),
        'lint': lint,
        'contact_validation': contact,
        'sensitive_data': sensitive,
        'timeline': timeline,
        'timeline_overlaps': overlaps,
        'timeline_durations': durations,
    }


def upload_job_description(
    *,
    jd_text: str,
    alias_graph: Optional[Dict[str, List[str]]] = None,
) -> Dict[str, Any]:
    parsed = parse_job_description(jd_text, alias_graph=alias_graph)
    graph = build_requirement_graph(parsed, alias_graph=alias_graph)
    return {'parsed_job': parsed, 'requirement_graph': graph}


def score_match(
    *,
    resume: CanonicalResume,
    jd_text: str,
    alias_graph: Optional[Dict[str, List[str]]] = None,
) -> Dict[str, Any]:
    score = compute_match_score(resume, jd_text, alias_graph=alias_graph)
    return {'score': score, 'explanations': generate_match_explanations(score)}
